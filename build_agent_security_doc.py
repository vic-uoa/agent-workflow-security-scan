from pathlib import Path
import re
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "outputs"
OUT.mkdir(exist_ok=True)

DOCX_PATH = OUT / "Agent安全预研简报.docx"
FLOW_PATH = OUT / "agent_flow.png"
SCOPE_PATH = OUT / "agent_security_scope.png"

FONT_PATH = Path(r"C:\Windows\Fonts\NotoSansSC-VF.ttf")
if not FONT_PATH.exists():
    FONT_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")

CJK_FONT = "Microsoft YaHei"
ACCENT = "2E74B5"
DARK = "0B2545"
MUTED = "555555"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "D9E2EC"


def load_font(size, bold=False):
    return ImageFont.truetype(str(FONT_PATH), size=size)


def wrap_zh(draw, text, font, max_width):
    lines = []
    for block in text.split("\n"):
        current = ""
        for ch in block:
            trial = current + ch
            if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
                current = trial
            else:
                if current:
                    lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines


def rounded_rect(draw, xy, radius, fill, outline=None, width=2):
    if isinstance(fill, str) and not fill.startswith("#"):
        fill = "#" + fill
    if isinstance(outline, str) and not outline.startswith("#"):
        outline = "#" + outline
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def draw_centered(draw, box, text, font, fill, line_gap=8):
    x1, y1, x2, y2 = box
    lines = wrap_zh(draw, text, font, x2 - x1 - 28)
    line_h = font.size + line_gap
    total_h = len(lines) * line_h - line_gap
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        w = draw.textbbox((0, 0), line, font=font)[2]
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=font, fill=fill)
        y += line_h


def make_flow_diagram():
    img = Image.new("RGB", (1800, 720), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(42)
    box_font = load_font(28)
    small_font = load_font(24)
    caption_font = load_font(22)

    draw.text((70, 44), "Agent 工作流与主要安全插入点", font=title_font, fill=(11, 37, 69))
    draw.text(
        (70, 100),
        "安全重点不只在模型回答，而在“目标-规划-工具-执行-反馈”整条链路。",
        font=caption_font,
        fill=(85, 85, 85),
    )

    boxes = [
        ((70, 190, 330, 360), "用户目标\n与策略约束", "#E8EEF5"),
        ((420, 190, 680, 360), "LLM 推理\n任务规划", "#F4F6F9"),
        ((770, 190, 1030, 360), "工具 / MCP\nRAG / 记忆", "#FFF4D6"),
        ((1120, 190, 1380, 360), "外部系统\n执行动作", "#FDECEC"),
        ((1470, 190, 1730, 360), "观察反馈\n审计记录", "#EAF7EA"),
    ]
    for box, label, color in boxes:
        rounded_rect(draw, box, 24, color, BORDER, 3)
        draw_centered(draw, box, label, box_font, (11, 37, 69))

    arrow_y = 275
    for x in [345, 695, 1045, 1395]:
        draw.line((x, arrow_y, x + 55, arrow_y), fill=(46, 116, 181), width=5)
        draw.polygon(
            [(x + 55, arrow_y), (x + 37, arrow_y - 13), (x + 37, arrow_y + 13)],
            fill=(46, 116, 181),
        )

    risk_boxes = [
        ((135, 420, 520, 525), "直接/间接提示注入\n伪造目标、覆盖规则"),
        ((570, 420, 955, 525), "工具投毒 / 工具影子\n描述、参数、返回值藏指令"),
        ((1005, 420, 1390, 525), "权限滥用 / 命令执行\n越权读写、外联、交易"),
        ((615, 570, 1185, 655), "防护点：最小权限、隔离沙箱、工具签名、人工确认、日志与回放"),
    ]
    for i, (box, label) in enumerate(risk_boxes):
        fill = "#FFF7ED" if i < 3 else "#E8EEF5"
        rounded_rect(draw, box, 18, fill, "#F0B27A", 2)
        draw_centered(draw, box, label, small_font, (120, 70, 0), line_gap=7)

    img.save(FLOW_PATH)


def make_scope_diagram():
    img = Image.new("RGB", (1800, 760), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(42)
    box_font = load_font(26)
    small_font = load_font(22)

    draw.text((70, 44), "Agent 安全涉及范围分布", font=title_font, fill=(11, 37, 69))
    draw.text(
        (70, 100),
        "从 LLM 原生风险向工具链、数据链、执行链和治理链扩展。",
        font=small_font,
        fill=(85, 85, 85),
    )

    layers = [
        ("模型与提示", "越狱、提示注入、幻觉、策略绕过", "#E8EEF5"),
        ("代理编排", "目标劫持、计划偏移、多步任务串联风险", "#F4F6F9"),
        ("工具 / MCP", "工具投毒、影子工具、rug pull、参数注入", "#FFF4D6"),
        ("数据与记忆", "RAG 污染、记忆投毒、敏感数据进入上下文", "#EAF7EA"),
        ("执行权限", "命令执行、文件/API/浏览器/金融动作越权", "#FDECEC"),
        ("供应链与治理", "Skill/插件/依赖风险、审计、合规、运营监控", "#F2F4F7"),
    ]
    x, y = 100, 170
    w, h = 760, 70
    for idx, (name, desc, color) in enumerate(layers):
        yy = y + idx * 84
        rounded_rect(draw, (x, yy, x + w, yy + h), 18, color, BORDER, 2)
        draw.text((x + 28, yy + 15), name, font=box_font, fill=(11, 37, 69))
        draw.text((x + 230, yy + 19), desc, font=small_font, fill=(50, 50, 50))

    controls = [
        ("识别", "资产清单、MCP/Skill/工具枚举"),
        ("评估", "基准测试、红队、扫描、风险评分"),
        ("控制", "权限边界、沙箱、签名、输入输出校验"),
        ("运营", "CI 门禁、监控、审计、事件响应"),
    ]
    cx, cy = 1030, 180
    for idx, (name, desc) in enumerate(controls):
        yy = cy + idx * 120
        rounded_rect(draw, (cx, yy, 1680, yy + 88), 18, "#FFFFFF", "#2E74B5", 3)
        draw.text((cx + 28, yy + 14), name, font=box_font, fill=(46, 116, 181))
        draw.text((cx + 130, yy + 22), desc, font=small_font, fill=(50, 50, 50))

    draw.line((900, 200, 990, 200), fill=(46, 116, 181), width=4)
    draw.polygon([(990, 200), (970, 188), (970, 212)], fill=(46, 116, 181))
    draw.text((910, 220), "落地成安全流程", font=small_font, fill=(85, 85, 85))
    img.save(SCOPE_PATH)


def set_run_font(run, font=CJK_FONT, size=None, color=None, bold=None):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    sec.header_distance = Inches(0.49)
    sec.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = CJK_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("222222")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 15, 7),
        ("Heading 2", 13, ACCENT, 10, 5),
        ("Heading 3", 11.5, DARK, 8, 4),
    ]:
        st = styles[name]
        st.font.name = CJK_FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        st = styles[name]
        st.font.name = CJK_FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.18


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Agent 与 Agent 安全预研简报")
    set_run_font(r, size=24, color=DARK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("基于 AgentDojo、ASB、PyRIT、garak、Snyk Agent Scan、MCP Scanner、agent-security-scanner-mcp 与 OWASP ASI 的公开资料梳理")
    set_run_font(r, size=10.5, color=MUTED)

    p = doc.add_paragraph()
    r = p.add_run("资料阅读日期：2026-07-09 | 形式：预研简报 | 语言：中文")
    set_run_font(r, size=9.5, color=MUTED)

    add_callout(
        doc,
        "一句话结论",
        "Agent 安全不是“多一个 LLM 安全分类”，而是 LLM 被赋予工具、记忆、权限和长期上下文后形成的系统安全问题。它的核心是：让会规划、会调用工具、会改变外部状态的 AI 系统在可控边界内行动。",
    )


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    set_cell_border(cell, "D9E2EC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + "：")
    set_run_font(r, size=10.5, color=DARK, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color="222222")
    doc.add_paragraph()


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9E2EC", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ["top", "left", "bottom", "right"]:
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def format_table(table, header=True):
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    set_run_font(run, size=9.3)
            if header and row_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, size=9.5, color=DARK, bold=True)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:eastAsia"), CJK_FONT)
    r_fonts.set(qn("w:ascii"), CJK_FONT)
    r_fonts.set(qn("w:hAnsi"), CJK_FONT)
    r_pr.append(r_fonts)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_para(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED)


def add_project_summary(doc):
    doc.add_heading("项目逐项预研摘要", level=1)
    summaries = [
        (
            "AgentDojo",
            "ETH Zurich 与 Invariant Labs 相关团队发布的动态评测环境，核心是评估 LLM Agent 在工具调用任务中面对提示注入攻击和防御策略时的表现。它不是普通提示集，而是把用户任务、工具、攻击注入、Agent pipeline 和防御模块组合成可复现实验。可参考点是：如何设计 agent 安全基准、如何拆分任务套件、如何比较防御策略，以及如何把“能否完成用户任务”和“是否被攻击带偏”同时纳入评估。",
        ),
        (
            "ASB",
            "Agent Security Bench 是 ICLR 2025 工作，目标是形式化并系统评测 LLM-based agents 的攻击与防御，覆盖学业咨询、心理咨询、投资、法律建议等 10 类场景。它把直接提示注入、观察提示注入、Plan-of-Thought 后门、记忆投毒等攻击放进统一框架。可参考点是：agent 攻击面可以沿用户查询、观察数据、系统提示、记忆检索、工具选择等路径拆解，适合用于威胁建模和评测指标设计。",
        ),
        (
            "PyRIT",
            "PyRIT 是 Microsoft 的 Python Risk Identification Tool for generative AI；原 Azure/PyRIT 仓库已提示迁移到 microsoft/PyRIT。它更偏红队编排框架，支持自动化和人工协同测试、多轮攻击策略、标准场景、可插拔目标、记忆和评分器。可参考点是：把安全测试工程化为“目标、攻击策略、转换器、评分器、结果存储”的流水线，适合作为企业内部 GenAI/Agent 红队平台底座。",
        ),
        (
            "garak",
            "garak 是 NVIDIA 维护的 LLM vulnerability scanner，定位类似 LLM 领域的 nmap/Metasploit：用 probes 和 detectors 扫描模型或聊天系统在提示注入、越狱、数据泄露、幻觉、错误信息、毒性输出等方面的弱点。它本身更偏模型和对话系统扫描，不完全等同 agent 专项评测。可参考点是：漏洞探针插件化、检测器分离、命令行批量扫描、报告输出，这些思路可迁移到 agent 安全检测。",
        ),
        (
            "Snyk Agent Scan",
            "Snyk Agent Scan 面向本机/企业环境中的 agent 组件资产发现与扫描，覆盖 MCP servers、skills、agent 配置等，检测工具投毒、工具影子、toxic flows、skill 中的提示注入、恶意代码、凭据处理和硬编码密钥。它特别提醒：扫描 MCP 配置可能会执行配置中的 stdio server 命令，因此需要交互确认或沙箱。可参考点是：agent 安全不只扫代码，还要扫本机已安装能力、MCP 配置和技能供应链。",
        ),
        (
            "MCP Scanner",
            "mcpscanner.cloud 是面向 MCP Server 的免费开源扫描与评分服务，支持仓库扫描、配置扫描、批量扫描、REST API、GitHub Actions、公开 leaderboard 和 OWASP MCP Top 10 映射。公开文档称其包含 122 条规则、15 类漏洞，覆盖工具投毒、提示注入、rug pull、跨源升级、凭据泄露、影子 MCP server 等。可参考点是：MCP Server 上线前做安全评分、CI 门禁和合规导出。",
        ),
        (
            "agent-security-scanner-mcp",
            "sinewaveai/agent-security-scanner-mcp 把安全扫描包装成 MCP/CLI 能力，面向 Claude Code、Cursor、Windsurf、Cline 等 AI 编码工作流，覆盖代码漏洞、agent prompt 检测、agent action 执行前检查、MCP server 扫描、skill 扫描、依赖/幻觉包检测、SBOM 和合规证据。可参考点是：把安全检查嵌进 agent 自己的工具链，让 agent 在写代码、执行命令、添加依赖、提交 PR 前先过安全门。",
        ),
        (
            "OWASP Agentic Security Initiative",
            "OWASP GenAI Security Project 下的 Agentic Security Initiative 提供面向自治 agent 和多步 AI 工作流的社区治理框架，包括 Agentic Applications Top 10、MCP Server 安全开发指南、第三方 MCP Server 使用指南、市场/治理报告等。它不是单个扫描器，而是行业共识和控制框架。可参考点是：统一风险语言、建立组织级治理清单、把研发、SecOps、合规和业务负责人放到同一张 agent 风险地图上。",
        ),
    ]
    for name, body in summaries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run(name + "：")
        set_run_font(r, bold=True, color=DARK)
        r = p.add_run(body)
        set_run_font(r)


def add_project_matrix(doc):
    doc.add_heading("项目定位矩阵", level=1)
    data = [
        ["类别", "代表项目", "主要解决的问题", "适合怎么用"],
        ["研究评测", "AgentDojo、ASB", "定义攻击/防御实验，衡量 agent 是否被带偏", "用于构建内部 benchmark、选择评测指标"],
        ["红队编排", "PyRIT、garak", "批量化探测模型/应用弱点，沉淀攻击策略和评分", "用于上线前红队、回归测试和模型比较"],
        ["组件扫描", "Snyk Agent Scan、MCP Scanner、agent-security-scanner-mcp", "发现 MCP、skills、依赖、配置和代码中的风险", "用于资产清点、CI 门禁、开发者本地检查"],
        ["治理框架", "OWASP ASI", "统一 Top 10、指南、控制点和组织语言", "用于制度、流程、审计和跨团队协作"],
    ]
    table = doc.add_table(rows=len(data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [1600, 2100, 3100, 2560])
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
    format_table(table)
    doc.add_paragraph()


def add_sources(doc):
    doc.add_heading("参考资料", level=1)
    sources = [
        ("AgentDojo GitHub", "https://github.com/ethz-spylab/agentdojo"),
        ("AgentDojo Documentation", "https://agentdojo.spylab.ai/"),
        ("ASB GitHub", "https://github.com/agiresearch/ASB"),
        ("ASB Paper", "https://arxiv.org/abs/2410.02644"),
        ("PyRIT GitHub", "https://github.com/microsoft/PyRIT"),
        ("PyRIT Documentation", "https://microsoft.github.io/PyRIT/"),
        ("garak GitHub", "https://github.com/NVIDIA/garak"),
        ("garak Documentation", "https://docs.garak.ai/"),
        ("Snyk Agent Scan", "https://github.com/snyk/agent-scan"),
        ("MCP Scanner", "https://mcpscanner.cloud/"),
        ("MCP Scanner Documentation", "https://mcpscanner.cloud/docs"),
        ("agent-security-scanner-mcp", "https://github.com/sinewaveai/agent-security-scanner-mcp"),
        ("OWASP Agentic Security Initiative", "https://genai.owasp.org/initiatives/agentic-security-initiative/"),
        ("OWASP Top 10 for Agentic Applications 2026", "https://genai.owasp.org/resource/owasp-top-10-for-agentic-applications-for-2026/"),
        ("OWASP Practical Guide for Secure MCP Server Development", "https://genai.owasp.org/resource/a-practical-guide-for-secure-mcp-server-development/"),
    ]
    for label, url in sources:
        p = doc.add_paragraph(style="List Bullet")
        add_hyperlink(p, label, url)
        r = p.add_run(" - " + url)
        set_run_font(r, size=9.3, color=MUTED)


def build_doc():
    make_flow_diagram()
    make_scope_diagram()

    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("1. Agent 的背景与含义", level=1)
    add_para(
        doc,
        "在 LLM 语境下，Agent 通常指一个围绕目标自主推进任务的系统：它接收用户目标，调用大模型进行理解和规划，再通过工具、MCP server、浏览器、数据库、RAG、记忆或代码执行环境与外部世界交互，并根据反馈继续迭代。普通聊天机器人主要生成文本；Agent 则可能读文件、查系统、调用 API、写代码、发消息、改配置甚至触发业务动作。",
    )
    add_para(
        doc,
        "因此，Agent 的关键不只是“模型聪不聪明”，而是模型、工具、权限、上下文、记忆、外部数据和执行环境如何被编排。只要它能改变外部状态，就需要把它当作一个自动化执行系统，而不是一个单纯的问答界面。",
    )
    doc.add_picture(str(FLOW_PATH), width=Inches(6.5))
    add_caption(doc, "图 1：Agent 从目标到执行的闭环及主要安全插入点")

    doc.add_heading("2. Agent 安全的背景：由 LLM 风险衍生", level=1)
    add_para(
        doc,
        "LLM 原生风险包括提示注入、越狱、幻觉、训练/上下文数据泄露、输出不可预测等。Agent 把这些风险放大，是因为 LLM 的输出不再只是文字建议，而会进入工具选择、参数生成、计划执行和长期记忆。当外部网页、邮件、文档、工单、代码注释或工具描述中混入恶意自然语言时，模型可能把这些内容误当成更高优先级指令。",
    )
    add_para(
        doc,
        "这也是 agent 安全区别于传统应用安全的地方：传统漏洞多来自代码逻辑或配置错误；agent 漏洞还可能来自自然语言、工具元数据、上下文排序、模型解释偏差和多工具组合。一个单独看似安全的“读取邮件”工具，和一个“发送 HTTP 请求”工具组合后，就可能形成从私有数据到外部网络的泄露路径。",
    )

    doc.add_heading("3. 具体涉及范围", level=1)
    add_para(
        doc,
        "Agent 安全的范围可以按链路分层理解：输入与提示、模型规划、工具/MCP、数据与记忆、执行权限、供应链、监控治理。每一层都有自己的风险，也会与其他层组合成更高危的跨层攻击。",
    )
    doc.add_picture(str(SCOPE_PATH), width=Inches(6.5))
    add_caption(doc, "图 2：Agent 安全从模型层扩展到工具、数据、权限和治理")
    doc.add_page_break()

    scope_rows = [
        ["层面", "典型风险", "控制重点"],
        ["输入与提示", "直接提示注入、间接提示注入、越狱、策略覆盖", "输入来源标注、内容隔离、指令/数据分离、输出校验"],
        ["规划与推理", "目标劫持、错误计划、绕过审批、多步任务偏移", "任务分解审计、策略约束、关键步骤人工确认"],
        ["工具 / MCP", "工具投毒、工具影子、rug pull、参数注入、跨 server 影响", "工具白名单、签名/哈希、最小权限、工具描述审查"],
        ["记忆 / RAG", "记忆投毒、检索污染、敏感数据进入上下文", "数据分级、检索过滤、记忆写入审批、过期与清理"],
        ["执行权限", "命令执行、文件改写、API 越权、财务或生产动作误触发", "沙箱、权限分层、审批门、回滚与审计"],
        ["供应链", "恶意 skill、插件、MCP server、幻觉包、依赖漏洞", "组件清单、来源验证、CI 扫描、版本锁定"],
        ["运营治理", "缺日志、难追责、无基线、无红队回归", "监控、证据留存、红队测试、合规映射"],
    ]
    table = doc.add_table(rows=len(scope_rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [1700, 3700, 3960])
    for r_idx, row in enumerate(scope_rows):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value
    format_table(table)

    doc.add_heading("4. 为什么现在有必要做 Agent 安全", level=1)
    for item in [
        "Agent 开始接入真实业务系统：从代码仓库、工单、邮件、知识库，到浏览器、终端、数据库和 SaaS API，错误或被操纵的动作会产生真实影响。",
        "MCP 等协议降低了工具接入门槛，也扩大了工具供应链：第三方 server、skills、插件和依赖需要像软件包一样被盘点、验证和持续监控。",
        "攻击入口更隐蔽：恶意指令可能藏在网页、文档、邮件、工具描述、返回值、记忆片段或依赖名称里，人工 review 很难只靠肉眼覆盖。",
        "传统 AppSec 控制不够：SAST/依赖扫描仍然必要，但还要增加提示注入评估、工具描述审查、agent 行为监控、权限边界和红队回归。",
    ]:
        add_para(doc, item, style="List Bullet")

    add_project_summary(doc)
    add_project_matrix(doc)

    doc.add_heading("建议的参考路线", level=1)
    add_para(
        doc,
        "若要在内部建立 agent 安全能力，建议按“资产识别 -> 威胁建模 -> 基准评测 -> 工程扫描 -> 运行时控制 -> 治理沉淀”推进。研究评测可参考 AgentDojo 与 ASB；红队编排可参考 PyRIT 与 garak；MCP/skill/依赖扫描可参考 Snyk Agent Scan、MCP Scanner 和 agent-security-scanner-mcp；组织语言和风险框架可参考 OWASP Agentic Security Initiative。",
    )
    add_callout(
        doc,
        "落地优先级",
        "先做资产清单和权限边界，再做提示注入/工具投毒专项评测；先把高危工具放进审批与沙箱，再逐步建设自动化扫描、CI 门禁和审计证据。",
    )
    add_sources(doc)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Agent 安全预研简报 | 2026-07-09")
    set_run_font(r, size=8.5, color=MUTED)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
