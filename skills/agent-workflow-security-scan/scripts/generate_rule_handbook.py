#!/usr/bin/env python3
"""Generate the audited security-rule handbook as a Word document."""

from __future__ import annotations

import argparse
import os
import re
import sys
from datetime import date
from pathlib import Path

import yaml


def _load_docx() -> None:
    try:
        import docx  # noqa: F401
        return
    except ModuleNotFoundError:
        runtime = Path(os.environ.get("USERPROFILE", "")) / ".cache/codex-runtimes/codex-primary-runtime/dependencies/python/Lib/site-packages"
        if runtime.exists():
            sys.path.append(str(runtime))
        import docx  # noqa: F401


_load_docx()

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
FAMILIES = ["FLOW", "IN", "LLM", "TOOL", "OUT", "KB"]
FAMILY_TITLES = {
    "FLOW": "工作流图与跨节点路径",
    "IN": "输入与 Start 节点",
    "LLM": "LLM 与 Agent 节点",
    "TOOL": "工具、HTTP 与代码执行",
    "OUT": "输出、End 与 Answer 节点",
    "KB": "知识库与持久记忆",
}
FAMILY_DESCRIPTIONS = {
    "FLOW": "验证跨节点可达性、分支、数据流与组合攻击链。",
    "IN": "验证输入契约、边界、敏感数据与高权限 Prompt 插值。",
    "LLM": "验证模型指令边界、结构化输出、授权委托与运行预算。",
    "TOOL": "验证工具能力、参数控制、网络外发、执行与对象级授权。",
    "OUT": "验证输出契约、受众、富文本渲染与高信任声明。",
    "KB": "验证知识范围、投毒链、来源治理、记忆隔离与持久化。",
}
DETECTABILITY = {"S": "静态", "H": "静态+运行时", "M": "人工/运行时"}
SEVERITY_COLORS = {
    "CRITICAL": "B42318",
    "HIGH": "C4320A",
    "MEDIUM": "B54708",
    "LOW": "175CD3",
    "INFO": "475467",
}


def parse_matrix(path: Path) -> dict[str, dict[str, str]]:
    rows: dict[str, dict[str, str]] = {}
    pattern = re.compile(
        r"^\| ((?:FLOW|IN|LLM|TOOL|OUT|KB)-\d{3}) \| (.*?) \| (.*?) \| (.*?) \|$",
        re.MULTILINE,
    )
    for match in pattern.finditer(path.read_text(encoding="utf-8")):
        rule_id, applicability, attack_chain, exclusion = match.groups()
        if rule_id in rows:
            raise ValueError(f"duplicate matrix row: {rule_id}")
        rows[rule_id] = {
            "applicability": applicability,
            "attack_chain": attack_chain,
            "exclusion": exclusion,
        }
    return rows


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=85, bottom=70, end=85) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_grid(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(Inches(width).emu / 635)))
        grid.append(col)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def set_run_font(run, size: float, *, bold: bool = False, color: str = "101828") -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_cell_text(cell, text: str, *, size: float = 8.0, bold: bool = False, color: str = "101828") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell)


def add_labeled_line(cell, label: str, value: str, *, size: float = 7.8) -> None:
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    set_run_font(paragraph.add_run(label), size, bold=True, color="344054")
    set_run_font(paragraph.add_run(value), size, color="101828")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(paragraph.add_run("内部安全规则手册  ·  "), 7.5, color="667085")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.extend([fld_char1, instr_text, fld_char2])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor.from_string("344054")
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for style_name, size, color in (
        ("Title", 26, "0B1F33"),
        ("Heading 1", 16, "0B1F33"),
        ("Heading 2", 12, "175CD3"),
    ):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.add_run("AGENT WORKFLOW SECURITY  /  RULEBOOK"), 7.2, bold=True, color="667085")
    add_page_number(section.footer.paragraphs[0])


def add_cover(doc: Document, rule_count: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    set_run_font(p.add_run("AGENT WORKFLOW SECURITY"), 10, bold=True, color="175CD3")
    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(12)
    title.add_run("安全规则细则手册")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(20)
    set_run_font(subtitle.add_run("Dify 工作流 DSL · 静态判定口径、排除条件与逐条示例"), 13, color="475467")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_repeat_table_grid(table, [3.2, 3.2, 3.2])
    values = [
        ("规则总数", str(rule_count)),
        ("规则分类", "6 类"),
        ("审查日期", date.today().isoformat()),
    ]
    for cell, (label, value) in zip(table.rows[0].cells, values):
        shade(cell, "F2F4F7")
        add_cell_text(cell, label, size=8, bold=True, color="667085")
        add_labeled_line(cell, "", value, size=15)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(24)
    note.paragraph_format.left_indent = Inches(0.05)
    set_run_font(note.add_run("版本口径："), 9, bold=True, color="344054")
    set_run_font(
        note.add_run("配置事实与模型可利用性分开评价；缺少运行时事实只记录覆盖缺口，不写成已确认漏洞。"),
        9,
        color="475467",
    )
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_audit_summary(doc: Document) -> None:
    doc.add_heading("审查结论", level=1)
    p = doc.add_paragraph()
    set_run_font(
        p.add_run(
            "规则库已按“适用能力 → 可达路径 → 缺失的匹配控制 → 业务影响 → 反证/降级”五步重新校准。"
            "规则目录、Dify 字段绑定、节点矩阵、引擎引用和逐条示例保持一一对应。"
        ),
        9.5,
    )

    changes = [
        ("Prompt 边界", "用户/知识内容插入 system 或 developer Prompt，作为 MEDIUM/CONFIRMED 配置事实；自动决策或高后果下游升为 HIGH。动态攻击是否成功单独验证。"),
        ("分支路由", "只有自由文本解析出的字段实际进入条件节点，且缺 enum、重复字段拒绝或失败关闭时才命中；不再仅凭 Code 节点存在推断风险。"),
        ("动态执行", "FLOW-010 仅覆盖变量可控代码体、命令或解释位置；固定代码把变量当数据解析不命中。"),
        ("多 Agent 通信", "FLOW-011 仅适用于 Agent/Agent-v2 之间的自由文本边界；普通 LLM 串联不再误报。"),
        ("高后果动作", "缺少不可绕过的确定性动作门可直接确认为 HIGH；动作门可以是策略、对象授权、参数约束或业务必要的人审。"),
        ("Markdown 外带", "静态目标可控仅定为 PROBABLE，必须核验渲染器是否自动取链、代理与 CSP 后再确认。"),
        ("部署与调用边界", "模型端点、调用方认证和下游消费方式不在 DSL 时只记 COVERAGE_GAP，不直接触发漏洞结论。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_repeat_table_grid(table, [2.0, 7.65])
    headers = ["优化主题", "校准后的生产判定"]
    for cell, value in zip(table.rows[0].cells, headers):
        shade(cell, "0B1F33")
        add_cell_text(cell, value, size=8.5, bold=True, color="FFFFFF")
    repeat_header(table.rows[0])
    for idx, (topic, decision) in enumerate(changes):
        row = table.add_row()
        keep_row_together(row)
        if idx % 2:
            for cell in row.cells:
                shade(cell, "F8FAFC")
        add_cell_text(row.cells[0], topic, size=8.2, bold=True, color="175CD3")
        add_cell_text(row.cells[1], decision, size=8.2)

    doc.add_heading("证据状态与严重性分离", level=2)
    state_table = doc.add_table(rows=1, cols=3)
    state_table.style = "Table Grid"
    set_repeat_table_grid(state_table, [1.55, 3.6, 4.5])
    for cell, value in zip(state_table.rows[0].cells, ["状态", "含义", "发布门禁作用"]):
        shade(cell, "EAF2FF")
        add_cell_text(cell, value, size=8.2, bold=True, color="175CD3")
    states = [
        ("CONFIRMED 已确认", "DSL、字段值、图路径或登记控制提供确定证据。", "MEDIUM 及以上进入 REVIEW；HIGH/CRITICAL 可阻断。"),
        ("PROBABLE 较可能", "核心路径存在，但影响仍依赖明确业务或运行时前提。", "进入 REVIEW，不写成已确认漏洞。"),
        ("OBSERVED 加固项", "确认了边界弱点，但尚未证明高后果能力或业务影响。", "LOW/INFO 单独不阻断发布。"),
        ("CANDIDATE 待验证", "只有分类候选或上下文线索，缺少完整资产/路径证据。", "进入人工复核，禁止升级为完整攻击链。"),
        ("COVERAGE_GAP 覆盖缺口", "所需事实不在导出 DSL 内。", "不作为漏洞计数；补充部署证据后重评。"),
        ("MITIGATED 已缓解", "风险路径存在，但被已验证、必经且覆盖所有路径的控制阻断。", "不阻断；保留控制证据。"),
    ]
    for state, meaning, gate in states:
        row = state_table.add_row()
        keep_row_together(row)
        add_cell_text(row.cells[0], state, size=7.8, bold=True)
        add_cell_text(row.cells[1], meaning, size=7.8)
        add_cell_text(row.cells[2], gate, size=7.8)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_rule_table(doc: Document, family: str, rules: list[dict], matrix: dict, examples: dict) -> None:
    doc.add_heading(f"{FAMILY_TITLES[family]}（{len(rules)} 条）", level=1)
    intro = doc.add_paragraph(FAMILY_DESCRIPTIONS[family])
    intro.paragraph_format.space_after = Pt(5)
    intro.paragraph_format.keep_with_next = True

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [1.45, 3.15, 2.05, 3.05]
    set_repeat_table_grid(table, widths)
    for cell, value in zip(table.rows[0].cells, ["规则 / 基准", "判定逻辑与攻击链", "排除或降级", "命中示例"]):
        shade(cell, "0B1F33")
        add_cell_text(cell, value, size=8.2, bold=True, color="FFFFFF")
        cell.paragraphs[0].paragraph_format.keep_with_next = True
    repeat_header(table.rows[0])

    for idx, rule in enumerate(rules):
        rule_id = rule["id"]
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        keep_row_together(row)
        if idx % 2:
            for cell in row.cells:
                shade(cell, "F8FAFC")

        first = row.cells[0]
        add_cell_text(first, rule_id, size=8.5, bold=True, color="175CD3")
        add_labeled_line(first, "", rule["title"], size=7.8)
        add_labeled_line(first, "基准：", rule["severity"], size=7.4)
        add_labeled_line(first, "检测：", DETECTABILITY.get(rule["detectability"], rule["detectability"]), size=7.4)
        add_labeled_line(first, "映射：", " / ".join(rule.get("standards", [])), size=7.0)
        for run in first.paragraphs[-3].runs:
            if rule["severity"] in SEVERITY_COLORS and run.text == rule["severity"]:
                run.font.color.rgb = RGBColor.from_string(SEVERITY_COLORS[rule["severity"]])
                run.font.bold = True

        logic = matrix[rule_id]
        add_cell_text(row.cells[1], logic["applicability"], size=7.8)
        add_labeled_line(row.cells[1], "攻击链：", logic["attack_chain"], size=7.6)
        add_cell_text(row.cells[2], logic["exclusion"], size=7.7)
        add_cell_text(row.cells[3], examples[rule_id], size=7.7)


def add_method_and_sources(doc: Document) -> None:
    doc.add_heading("使用方法与外部依据", level=1)
    doc.add_heading("逐条判定顺序", level=2)
    steps = [
        "确认规则所针对的能力真实存在，不能只靠节点标题或关键词。",
        "确认不可信/敏感源与目标节点之间存在具体变量、数据或控制流路径。",
        "确认缺少与该风险匹配且不可绕过的确定性控制。",
        "按业务消费方式和高后果能力确定严重性，不把加固建议写成漏洞。",
        "检查反证、平台控制与运行时缺失上下文，必要时降为较可能、待验证或覆盖缺口。",
    ]
    for index, item in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        set_run_font(p.add_run(f"{index}. "), 9, bold=True, color="175CD3")
        set_run_font(p.add_run(item), 9)

    doc.add_heading("主要依据", level=2)
    sources = [
        ("OWASP AISVS", "控制库存用于校准身份、工具、数据、记忆、输出和部署边界。", "https://github.com/OWASP/AISVS/blob/main/1.0/en/0x91-Appendix-B_AI_Security_Controls_Inventory.md"),
        ("OWASP LLM01 Prompt Injection", "支持将 Prompt 注入视为持续测试问题，并区分静态前置条件与模型可利用性。", "https://genai.owasp.org/llmrisk/llm01-prompt-injection/"),
        ("OpenAI Instruction Hierarchy", "支持高权限指令与低权限、不可信数据之间的层级隔离原则。", "https://openai.com/index/the-instruction-hierarchy/"),
        ("Greshake et al.", "说明间接 Prompt Injection 可通过外部内容影响集成应用。", "https://arxiv.org/abs/2302.12173"),
        ("NVIDIA garak", "用于核对模型层动态探测类别；不替代 DSL 静态证据。", "https://github.com/NVIDIA/garak"),
        ("promptfoo red-team strategies", "用于设计动态变形与组合测试，不允许模型结果反向覆盖确定性 Finding。", "https://github.com/promptfoo/promptfoo/blob/main/site/docs/red-team/strategies/index.md"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_repeat_table_grid(table, [2.0, 3.55, 4.1])
    for cell, value in zip(table.rows[0].cells, ["来源", "用于本规则库的部分", "链接"]):
        shade(cell, "0B1F33")
        add_cell_text(cell, value, size=8.2, bold=True, color="FFFFFF")
    repeat_header(table.rows[0])
    for source, use, url in sources:
        row = table.add_row()
        keep_row_together(row)
        add_cell_text(row.cells[0], source, size=8, bold=True)
        add_cell_text(row.cells[1], use, size=8)
        add_cell_text(row.cells[2], url, size=7.2, color="175CD3")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    set_run_font(p.add_run("限制："), 8.5, bold=True, color="B54708")
    set_run_font(
        p.add_run("本手册解释确定性规则。实际模型服从攻击指令、网络可达性、IAM、插件实现、渲染器行为与沙盒隔离仍需在受控环境动态验证。"),
        8.5,
    )


def build_document(output_path: Path) -> None:
    catalog = yaml.safe_load((ROOT / "rules/core-rules.yml").read_text(encoding="utf-8"))
    rules = catalog["rules"]
    matrix = parse_matrix(ROOT / "references/node-rule-matrix.md")
    examples = yaml.safe_load((ROOT / "references/rule-examples.yml").read_text(encoding="utf-8"))["examples"]
    rule_ids = [rule["id"] for rule in rules]
    if set(rule_ids) != set(matrix) or set(rule_ids) != set(examples):
        raise ValueError("catalog, matrix and examples must cover the same rule IDs")

    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "Agent Workflow 安全规则细则手册"
    doc.core_properties.subject = "Dify 工作流 DSL 静态安全规则"
    doc.core_properties.author = "Agent Workflow Security"
    add_cover(doc, len(rules))
    add_audit_summary(doc)
    for family in FAMILIES:
        add_rule_table(doc, family, [r for r in rules if r["id"].startswith(family + "-")], matrix, examples)
    add_method_and_sources(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output",
        type=Path,
        default=ROOT.parents[1] / "outputs/规则库审查/Agent-Workflow-安全规则细则手册.docx",
    )
    args = parser.parse_args()
    build_document(args.output.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
